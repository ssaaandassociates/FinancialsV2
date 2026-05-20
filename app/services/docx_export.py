"""
Word (.docx) Export Service
Generates formatted Financial Statement document with:
- Balance Sheet, Profit & Loss, Notes
- Signing block
- Professional formatting with A4 pages
"""
import os
from sqlalchemy.orm import Session
from app.models import Project, Client, SigningBlock
from app.services import financial_engine, notes_engine
from app.config import OUTPUT_DIR


def generate_docx(db: Session, project_id: int) -> str:
    """Generate Word document of financial statements."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise ValueError(f"Project {project_id} not found")
    client = project.client
    company = client.name

    bs = financial_engine.generate_bs(db, project_id)
    pl = financial_engine.generate_pl(db, project_id)
    notes = notes_engine.generate_all_notes(db, project_id)
    sb = db.query(SigningBlock).filter(SigningBlock.project_id == project_id).first()

    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    navy = RGBColor(0x1a, 0x27, 0x44)
    gold = RGBColor(0xc8, 0x97, 0x3e)

    def add_title(text, size=14):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = navy

    def add_subtitle(text, size=10):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x2d, 0x3f, 0x5e)

    def fmt(val):
        if val is None or val == 0:
            return "-"
        return f"{val:,.0f}"

    def add_fs_table(headers, rows):
        """Add a formatted financial statement table."""
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        # Header row
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(8)

        # Data rows
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = str(val) if val else ""
                for p in cell.paragraphs:
                    if c_idx >= 2:
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in p.runs:
                        run.font.size = Pt(8)
                        # Bold total rows
                        if isinstance(row[1], str) and ('TOTAL' in row[1].upper() or 'PROFIT' in row[1].upper()):
                            run.bold = True

    def add_signing():
        if not sb:
            return
        doc.add_paragraph()
        p = doc.add_paragraph("As per our report of even date attached")
        p.runs[0].font.size = Pt(9)
        doc.add_paragraph()

        table = doc.add_table(rows=6, cols=2)
        cells = [
            (0, 0, f"For {sb.auditor_firm or ''}"),
            (0, 1, f"For and on behalf of Board of Directors of {company}"),
            (1, 0, f"Chartered Accountants, FRN: {sb.auditor_frn or ''}"),
            (1, 1, ""),
            (2, 0, f"{sb.partner_name or ''}\nPartner, M.No.: {sb.partner_membership_no or ''}"),
            (2, 1, f"{sb.director1_name or ''} ({sb.director1_designation or 'Director'}), DIN: {sb.director1_din or ''}\n{sb.director2_name or ''} ({sb.director2_designation or 'Director'}), DIN: {sb.director2_din or ''}"),
            (3, 0, f"Place: {sb.place or ''}"),
            (3, 1, f"Place: {sb.place or ''}"),
            (4, 0, f"Date: {sb.signing_date or ''}"),
            (4, 1, f"Date: {sb.signing_date or ''}"),
            (5, 0, f"UDIN: {sb.partner_udin or ''}"),
            (5, 1, ""),
        ]
        for r, c, text in cells:
            table.rows[r].cells[c].text = text
            for p in table.rows[r].cells[c].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    # ===== BALANCE SHEET =====
    add_title(company)
    add_subtitle(f"Balance Sheet {project.bs_header_cy}")

    bs_rows = []
    for item in bs["line_items"]:
        sno, text, note_ref, cy, py, level, is_total = item
        bs_rows.append([sno or "", text or "", note_ref or "", fmt(cy), fmt(py)])

    add_fs_table(
        ["", "Particulars", "Note", project.bs_header_cy, project.bs_header_py],
        bs_rows
    )
    add_signing()
    doc.add_page_break()

    # ===== PROFIT & LOSS =====
    add_title(company)
    add_subtitle(f"Statement of Profit and Loss {project.pl_header_cy}")

    pl_rows = []
    for item in pl["line_items"]:
        sno, text, note_ref, cy, py, level, is_total = item
        pl_rows.append([sno or "", text or "", note_ref or "", fmt(cy), fmt(py)])

    add_fs_table(
        ["", "Particulars", "Note", project.pl_header_cy, project.pl_header_py],
        pl_rows
    )
    add_signing()
    doc.add_page_break()

    # ===== NOTES =====
    add_title(company)
    add_subtitle("Notes to the Financial Statements")

    for note_key, note_data in notes["bs_notes"].items():
        p = doc.add_paragraph()
        run = p.add_run(f"Note {note_key}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = navy

        for item in note_data["items"]:
            p = doc.add_paragraph()
            run = p.add_run(f"  {item['sno']}. {item['particulars']}")
            run.font.size = Pt(9)
            run2 = p.add_run(f"    CY: {fmt(item['cy_amount'])}    PY: {fmt(item['py_amount'])}")
            run2.font.size = Pt(9)

    # Save
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    filename = f"FS_{client.name[:20]}_{project.financial_year}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename.replace(" ", "_"))
    doc.save(filepath)
    return filepath
